import os
import re
import base64
import html as html_lib
from html.parser import HTMLParser
import ssl
import urllib.request
import zipfile
from io import BytesIO
from datetime import datetime, date
from functools import wraps
import psycopg2
import psycopg2.extras
from fpdf import FPDF
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask import (
    Flask, render_template, redirect, url_for,
    request, session, flash, g, make_response
)
import cloudinary
import cloudinary.uploader
from dotenv import load_dotenv

load_dotenv()

# ── Cloudinary ────────────────────────────────────────────────────────────────
cloudinary.config(
    cloud_name=os.environ.get("CLOUDINARY_CLOUD_NAME"),
    api_key=os.environ.get("CLOUDINARY_API_KEY"),
    api_secret=os.environ.get("CLOUDINARY_API_SECRET"),
    secure=True,
)

# ── PDF helpers ───────────────────────────────────────────────────────────────

# DejaVu Sans font paths (downloaded on first PDF generation)
_FONTS_DIR = os.path.join(os.path.dirname(__file__), 'static', 'fonts')
_DEJAVU_ZIP  = 'https://github.com/dejavu-fonts/dejavu-fonts/releases/download/version_2_37/dejavu-fonts-ttf-2.37.zip'
_DEJAVU_FILES = ['DejaVuSans.ttf', 'DejaVuSans-Bold.ttf', 'DejaVuSans-Oblique.ttf', 'DejaVuSans-BoldOblique.ttf']

def _ensure_fonts():
    """Download DejaVu Sans TTF files if not already cached (downloads release zip once)."""
    os.makedirs(_FONTS_DIR, exist_ok=True)
    needed = [f for f in _DEJAVU_FILES if not os.path.exists(os.path.join(_FONTS_DIR, f))]
    if not needed:
        return
    # Bypass SSL cert check — needed on macOS Python installs without certs
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    req = urllib.request.Request(_DEJAVU_ZIP, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req, context=ctx) as r:
        zip_bytes = r.read()
    with zipfile.ZipFile(BytesIO(zip_bytes)) as zf:
        for fname in needed:
            inner = f'dejavu-fonts-ttf-2.37/ttf/{fname}'
            with zf.open(inner) as src, open(os.path.join(_FONTS_DIR, fname), 'wb') as dst:
                dst.write(src.read())

def _safe(text):
    """Minimal text clean-up for Unicode font output (no encoding restrictions)."""
    return text.replace('\u00a0', ' ') if text else ''  # just strip non-breaking spaces


class _HTMLToPDF(HTMLParser):
    """Walk HTML and write it into an FPDF doc preserving bold, italic, headings, lists."""

    def __init__(self, pdf, base_size=11, lh=6.5):
        super().__init__()
        self.pdf = pdf
        self.base_size = base_size
        self.lh = lh
        self.bold = 0
        self.italic = 0
        self.skip = 0
        self.list_stack = []   # [['ul'|'ol', counter], ...]
        self.at_bol = True
        self._orig_lm = pdf.l_margin   # page left margin, restored after each li
        self._font(base_size)

    def _font(self, size=None):
        style = ('B' if self.bold else '') + ('I' if self.italic else '')
        self.pdf.set_font('DejaVu', style, size or self.base_size)

    def _newline(self, gap=0):
        if not self.at_bol:
            self.pdf.ln(self.lh)
        if gap:
            self.pdf.ln(gap)
        self.pdf.set_x(self.pdf.l_margin)
        self.at_bol = True

    def _write(self, text):
        text = _safe(text)
        if not text:
            return
        self._font()
        self.pdf.set_x(self.pdf.l_margin)
        self.pdf.multi_cell(0, self.lh, text, align="J")
        self.at_bol = True

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in ('script', 'style'):
            self.skip += 1; return
        if self.skip:
            return
        if tag in ('b', 'strong'):
            self.bold += 1; self._font()
        elif tag in ('i', 'em'):
            self.italic += 1; self._font()
        elif tag == 'h2':
            self._newline(3); self.bold += 1; self._font(13)
        elif tag == 'h3':
            self._newline(2); self.bold += 1; self._font(12)
        elif tag in ('p', 'div'):
            self._newline(2)
        elif tag == 'br':
            self._newline()
        elif tag == 'ul':
            self._newline(); self.list_stack.append(['ul', 0])
        elif tag == 'ol':
            self._newline(); self.list_stack.append(['ol', 0])
        elif tag == 'li':
            self._newline()
            self.pdf.set_left_margin(self._orig_lm)
            depth = len(self.list_stack)
            indent = depth * 5        # mm indent per nesting level
            bullet_w = 5              # mm reserved for bullet + gap
            if self.list_stack:
                lst = self.list_stack[-1]
                if lst[0] == 'ol':
                    lst[1] += 1; bullet = f"{lst[1]}."
                else:
                    bullet = '•'  # circle bullet (cp1252 0x95)
            else:
                bullet = '•'
            x_bullet = self.pdf.l_margin + indent
            x_text   = x_bullet + bullet_w
            # Write bullet, then set hanging indent for wrapped lines
            self.pdf.set_x(x_bullet)
            self._font()
            self.pdf.cell(bullet_w, self.lh, _safe(bullet))
            self.pdf.set_left_margin(x_text)
            self.pdf.set_x(x_text)
            self.at_bol = False

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ('script', 'style'):
            self.skip = max(0, self.skip - 1); return
        if self.skip:
            return
        if tag in ('b', 'strong'):
            self.bold = max(0, self.bold - 1); self._font()
        elif tag in ('i', 'em'):
            self.italic = max(0, self.italic - 1); self._font()
        elif tag in ('h2', 'h3'):
            self.bold = max(0, self.bold - 1); self._newline(3); self._font(self.base_size)
        elif tag in ('p', 'div'):
            self._newline(2)
        elif tag in ('ul', 'ol'):
            if self.list_stack: self.list_stack.pop()
            self.pdf.set_left_margin(self._orig_lm)
            self._newline(1)
        elif tag == 'li':
            self._newline()
            self.pdf.set_left_margin(self._orig_lm)

    def handle_data(self, data):
        if self.skip:
            return
        text = re.sub(r'\s+', ' ', data)
        if text.strip():
            self._write(text)


def _render_html(pdf, html_content, base_size=11, lh=6.5):
    """Render HTML into the FPDF document at the current cursor position."""
    pdf.set_x(pdf.l_margin)
    r = _HTMLToPDF(pdf, base_size, lh)
    r.feed(html_lib.unescape(html_content or ''))
    if not r.at_bol:
        pdf.ln(lh)
    pdf.set_x(pdf.l_margin)
    pdf.set_font('DejaVu', '', base_size)


# ── Word document helpers ─────────────────────────────────────────────────────

class _HTMLToDocx(HTMLParser):
    """Parse HTML content and write it into a python-docx Document."""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.bold = 0
        self.italic = 0
        self.skip = 0
        self.list_stack = []   # [['ul'|'ol', counter], ...]
        self.current_para = None
        self._heading_size = None  # Pt size when inside an h2/h3

    def _end_para(self):
        self.current_para = None
        self._heading_size = None

    def _get_para(self):
        if self.current_para is None:
            self.current_para = self.doc.add_paragraph(style='Normal')
            self.current_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self.current_para.paragraph_format.space_after = Pt(8)
        return self.current_para

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in ('script', 'style'):
            self.skip += 1; return
        if self.skip: return
        if tag in ('b', 'strong'):
            self.bold += 1
        elif tag in ('i', 'em'):
            self.italic += 1
        elif tag == 'h2':
            self._end_para()
            self.current_para = self.doc.add_paragraph(style='Normal')
            self.current_para.paragraph_format.space_before = Pt(10)
            self.current_para.paragraph_format.space_after  = Pt(4)
            self.current_para.paragraph_format.keep_with_next = True
            self._heading_size = 13
            self.bold += 1
        elif tag == 'h3':
            self._end_para()
            self.current_para = self.doc.add_paragraph(style='Normal')
            self.current_para.paragraph_format.space_before = Pt(8)
            self.current_para.paragraph_format.space_after  = Pt(4)
            self.current_para.paragraph_format.keep_with_next = True
            self._heading_size = 12
            self.bold += 1
        elif tag == 'p':
            self._end_para()
        elif tag == 'br':
            self._get_para().add_run().add_break()
        elif tag == 'ul':
            self._end_para()
            self.list_stack.append(['ul', 0])
        elif tag == 'ol':
            self._end_para()
            self.list_stack.append(['ol', 0])
        elif tag == 'li':
            self._end_para()
            if self.list_stack and self.list_stack[-1][0] == 'ol':
                self.list_stack[-1][1] += 1
                self.current_para = self.doc.add_paragraph(style='List Number')
            else:
                self.current_para = self.doc.add_paragraph(style='List Bullet')

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in ('script', 'style'):
            self.skip = max(0, self.skip - 1); return
        if self.skip: return
        if tag in ('b', 'strong'):
            self.bold = max(0, self.bold - 1)
        elif tag in ('i', 'em'):
            self.italic = max(0, self.italic - 1)
        elif tag in ('h2', 'h3'):
            self.bold = max(0, self.bold - 1)
            self._end_para()
        elif tag in ('p', 'li'):
            self._end_para()
        elif tag in ('ul', 'ol'):
            if self.list_stack: self.list_stack.pop()
            self._end_para()

    def handle_data(self, data):
        if self.skip: return
        text = re.sub(r'\s+', ' ', data)
        if not text.strip(): return
        run = self._get_para().add_run(text)
        run.bold = self.bold > 0
        run.italic = self.italic > 0
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(self._heading_size if self._heading_size else 11)


def _render_html_docx(doc, html_content):
    """Render HTML content into a python-docx Document."""
    renderer = _HTMLToDocx(doc)
    renderer.feed(html_lib.unescape(html_content or ''))


app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "change-this-secret-key")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "music2024")

# ── Database ──────────────────────────────────────────────────────────────────

def get_db():
    if "db" not in g:
        url = os.environ.get("DATABASE_URL", "")
        if url.startswith("postgres://"):
            url = url.replace("postgres://", "postgresql://", 1)
        g.db = psycopg2.connect(url)
    return g.db

@app.teardown_appcontext
def close_db(error):
    db = g.pop("db", None)
    if db is not None:
        try:
            if error:
                db.rollback()
            else:
                db.commit()
        finally:
            db.close()

def _swap_sort_order(table, id_col, id1, id2):
    conn = get_db()
    with conn.cursor() as c:
        c.execute(f"SELECT sort_order FROM {table} WHERE {id_col} = %s", (id1,))
        order1 = c.fetchone()[0]
        c.execute(f"SELECT sort_order FROM {table} WHERE {id_col} = %s", (id2,))
        order2 = c.fetchone()[0]
        c.execute(f"UPDATE {table} SET sort_order = %s WHERE {id_col} = %s", (order2, id1))
        c.execute(f"UPDATE {table} SET sort_order = %s WHERE {id_col} = %s", (order1, id2))

def new_id():
    return int(datetime.now().timestamp())

def get_max_sort_order(table, where_col=None, where_val=None):
    with get_db().cursor() as c:
        if where_col:
            c.execute(f"SELECT COALESCE(MAX(sort_order), -1) FROM {table} WHERE {where_col} = %s", (where_val,))
        else:
            c.execute(f"SELECT COALESCE(MAX(sort_order), -1) FROM {table}")
        return c.fetchone()[0]

# ── Data loaders ──────────────────────────────────────────────────────────────

def load_announcements():
    with get_db().cursor(cursor_factory=psycopg2.extras.RealDictCursor) as c:
        c.execute("SELECT * FROM announcements ORDER BY date DESC, id DESC")
        return [dict(r) for r in c.fetchall()]

def load_handbook():
    """Full load including chapter content — only use when content is actually needed."""
    with get_db().cursor(cursor_factory=psycopg2.extras.RealDictCursor) as c:
        c.execute("SELECT * FROM handbook_sections ORDER BY sort_order, id")
        sections = [dict(r) for r in c.fetchall()]
        for section in sections:
            c.execute(
                "SELECT * FROM handbook_chapters WHERE section_id = %s ORDER BY sort_order, id",
                (section["id"],)
            )
            section["chapters"] = [dict(r) for r in c.fetchall()]
    return sections

def load_handbook_meta():
    """Lightweight load — sections with chapter titles only, no content. Use for nav/listings."""
    with get_db().cursor(cursor_factory=psycopg2.extras.RealDictCursor) as c:
        c.execute("SELECT id, title, description, sort_order FROM handbook_sections ORDER BY sort_order, id")
        sections = [dict(r) for r in c.fetchall()]
        for section in sections:
            c.execute(
                "SELECT id, title, section_id, sort_order FROM handbook_chapters WHERE section_id = %s ORDER BY sort_order, id",
                (section["id"],)
            )
            section["chapters"] = [dict(r) for r in c.fetchall()]
    return sections

def load_resources():
    with get_db().cursor(cursor_factory=psycopg2.extras.RealDictCursor) as c:
        c.execute("SELECT * FROM resource_categories ORDER BY sort_order, id")
        categories = [dict(r) for r in c.fetchall()]
        for cat in categories:
            c.execute(
                "SELECT * FROM resource_items WHERE category_id = %s ORDER BY sort_order, id",
                (cat["id"],)
            )
            cat["items"] = [dict(r) for r in c.fetchall()]
    return categories

def get_category_id_by_name(name):
    with get_db().cursor() as c:
        c.execute("SELECT id FROM resource_categories WHERE name = %s", (name,))
        row = c.fetchone()
        return row[0] if row else None

# ── Auth ──────────────────────────────────────────────────────────────────────

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("admin"):
            return redirect(url_for("admin_login"))
        return f(*args, **kwargs)
    return decorated

# ── Health check (used by UptimeRobot to keep the app awake) ──────────────────

@app.route("/ping")
def ping():
    return "OK", 200

# ── Public routes ─────────────────────────────────────────────────────────────

@app.route("/")
def index():
    with get_db().cursor(cursor_factory=psycopg2.extras.RealDictCursor) as c:
        c.execute("SELECT id, title, date, tag FROM announcements ORDER BY date DESC, id DESC LIMIT 3")
        recent = [dict(r) for r in c.fetchall()]
    return render_template("index.html", announcements=recent)

@app.route("/announcements")
def announcements():
    items = load_announcements()
    return render_template("announcements.html", announcements=items)

@app.route("/handbook")
def handbook():
    # Only needs section titles/descriptions for the bento cards — no chapter content
    sections = load_handbook_meta()
    return render_template("handbook.html", sections=sections)

@app.route("/handbook/<int:section_id>")
def handbook_section(section_id):
    # Sidebar needs section titles; main content needs chapter content for this section only
    all_sections = load_handbook_meta()
    section = next((s for s in all_sections if s["id"] == section_id), None)
    if not section:
        flash("Section not found.", "error")
        return redirect(url_for("handbook"))
    with get_db().cursor(cursor_factory=psycopg2.extras.RealDictCursor) as c:
        c.execute(
            "SELECT id, title, content, sort_order FROM handbook_chapters WHERE section_id = %s ORDER BY sort_order, id",
            (section_id,)
        )
        section["chapters"] = [dict(r) for r in c.fetchall()]
    return render_template("handbook_section.html", section=section, sections=all_sections)

@app.route("/handbook/<int:section_id>/chapter/<int:chapter_id>")
def handbook_chapter(section_id, chapter_id):
    # Fetch only what this page needs: section title, chapter titles for prev/next, and this chapter's content
    with get_db().cursor(cursor_factory=psycopg2.extras.RealDictCursor) as c:
        c.execute("SELECT id, title FROM handbook_sections WHERE id = %s", (section_id,))
        section_row = c.fetchone()
        if not section_row:
            return redirect(url_for("handbook"))
        section = dict(section_row)
        c.execute(
            "SELECT id, title FROM handbook_chapters WHERE section_id = %s ORDER BY sort_order, id",
            (section_id,)
        )
        chapters = [dict(r) for r in c.fetchall()]
        chapter = next((ch for ch in chapters if ch["id"] == chapter_id), None)
        if not chapter:
            return redirect(url_for("handbook_section", section_id=section_id))
        c.execute("SELECT content FROM handbook_chapters WHERE id = %s", (chapter_id,))
        content_row = c.fetchone()
        if content_row:
            chapter["content"] = content_row["content"]
    idx = next(i for i, ch in enumerate(chapters) if ch["id"] == chapter_id)
    prev_ch = chapters[idx - 1] if idx > 0 else None
    next_ch = chapters[idx + 1] if idx < len(chapters) - 1 else None
    return render_template("handbook_chapter.html",
        section=section, chapter=chapter,
        prev_ch=prev_ch, next_ch=next_ch)

@app.route("/resources")
def resources():
    categories = load_resources()
    return render_template("resources.html", categories=categories)

@app.route("/resources/article/<int:item_id>")
def resource_article(item_id):
    categories = load_resources()
    for cat in categories:
        for item in cat["items"]:
            if item["id"] == item_id:
                siblings = [i for i in cat["items"] if i.get("type") == "article"]
                return render_template("resource_article.html",
                    item=item, category_name=cat["name"], siblings=siblings)
    flash("Article not found.", "error")
    return redirect(url_for("resources"))

# ── Admin: login ──────────────────────────────────────────────────────────────

@app.route("/admin/login", methods=["GET", "POST"])
def admin_login():
    if request.method == "POST":
        if request.form.get("password") == ADMIN_PASSWORD:
            session["admin"] = True
            return redirect(url_for("admin_dashboard"))
        flash("Incorrect password.", "error")
    return render_template("admin/login.html")

@app.route("/admin/logout")
def admin_logout():
    session.pop("admin", None)
    return redirect(url_for("index"))

# ── Admin: dashboard ──────────────────────────────────────────────────────────

@app.route("/admin")
@login_required
def admin_dashboard():
    with get_db().cursor() as c:
        c.execute("SELECT COUNT(*) FROM announcements")
        announcement_count = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM handbook_chapters")
        handbook_count = c.fetchone()[0]
        c.execute("SELECT COUNT(*) FROM resource_items")
        resource_count = c.fetchone()[0]
    return render_template("admin/dashboard.html",
        announcement_count=announcement_count,
        handbook_count=handbook_count,
        resource_count=resource_count,
    )

# ── Admin: announcements ──────────────────────────────────────────────────────

@app.route("/admin/announcements")
@login_required
def admin_announcements():
    return render_template("admin/announcements.html", announcements=load_announcements())

def _extract_image(request):
    """Upload image to Cloudinary and return the URL, or fall back to a pasted URL."""
    file = request.files.get("image_file")
    if file and file.filename:
        data = file.read()
        result = cloudinary.uploader.upload(
            data,
            folder="cos-mm",
            resource_type="image",
        )
        return result["secure_url"]
    return request.form.get("image_url", "")

@app.route("/admin/announcements/new", methods=["GET", "POST"])
@login_required
def admin_announcement_new():
    if request.method == "POST":
        image_url = _extract_image(request)
        with get_db().cursor() as c:
            c.execute(
                "INSERT INTO announcements (id, title, body, date, tag, image_url) VALUES (%s, %s, %s, %s, %s, %s)",
                (
                    new_id(),
                    request.form["title"],
                    request.form["body"],
                    request.form["date"] or datetime.today().strftime("%Y-%m-%d"),
                    request.form.get("tag", ""),
                    image_url,
                )
            )
        flash("Announcement added.", "success")
        return redirect(url_for("admin_announcements"))
    return render_template("admin/announcement_form.html", item=None)

@app.route("/admin/announcements/edit/<int:item_id>", methods=["GET", "POST"])
@login_required
def admin_announcement_edit(item_id):
    with get_db().cursor(cursor_factory=psycopg2.extras.RealDictCursor) as c:
        c.execute("SELECT * FROM announcements WHERE id = %s", (item_id,))
        item = c.fetchone()
    if not item:
        flash("Not found.", "error")
        return redirect(url_for("admin_announcements"))
    item = dict(item)
    if request.method == "POST":
        image_url = _extract_image(request)
        # Keep existing image if no new one provided
        if not image_url:
            image_url = item.get("image_url", "")
        with get_db().cursor() as c:
            c.execute(
                "UPDATE announcements SET title=%s, body=%s, date=%s, tag=%s, image_url=%s WHERE id=%s",
                (request.form["title"], request.form["body"],
                 request.form["date"], request.form.get("tag", ""), image_url, item_id)
            )
        flash("Announcement updated.", "success")
        return redirect(url_for("admin_announcements"))
    return render_template("admin/announcement_form.html", item=item)

@app.route("/admin/announcements/delete/<int:item_id>", methods=["POST"])
@login_required
def admin_announcement_delete(item_id):
    with get_db().cursor() as c:
        c.execute("DELETE FROM announcements WHERE id = %s", (item_id,))
    flash("Announcement deleted.", "success")
    return redirect(url_for("admin_announcements"))

# ── Admin: handbook sections ──────────────────────────────────────────────────

@app.route("/admin/handbook")
@login_required
def admin_handbook():
    return render_template("admin/handbook.html", sections=load_handbook_meta())

@app.route("/admin/handbook/docx")
@login_required
def admin_handbook_docx():
    sections = load_handbook()

    doc = DocxDocument()

    # Page margins (1.25" left/right, 1" top/bottom)
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    def _h(text, size, space_before=0, space_after=8, center=False, bold=True):
        """Plain black bold paragraph — avoids Word's blue themed heading styles."""
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_before  = Pt(space_before)
        p.paragraph_format.space_after   = Pt(space_after)
        p.paragraph_format.keep_with_next = not center
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    # ── Cover page ────────────────────────────────────────────────────────────
    # space_before pushes title ~⅓ down the page body (9" body ≈ 648pt; ⅓ ≈ 216pt)
    _h('Co-workers Handbook', 32, space_before=180, space_after=14, center=True)

    sub = doc.add_paragraph(style='Normal')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after  = Pt(6)
    sr = sub.add_run('Church of Singapore Music Ministry')
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0, 0, 0)

    dt = doc.add_paragraph(style='Normal')
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dtr = dt.add_run(date.today().strftime('%B %Y'))
    dtr.font.size = Pt(11)
    dtr.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()

    # ── Table of Contents ─────────────────────────────────────────────────────
    _h('Contents', 22, space_before=0, space_after=14)

    for i, section in enumerate(sections, 1):
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"{i}.  {section['title']}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)

        for j, chapter in enumerate(section.get('chapters', []), 1):
            cp = doc.add_paragraph(style='Normal')
            cp.paragraph_format.left_indent  = Inches(0.3)
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after  = Pt(2)
            cr = cp.add_run(f"{i}.{j}  {chapter['title']}")
            cr.font.size = Pt(10)
            cr.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()

    # ── Sections & Chapters ───────────────────────────────────────────────────
    for i, section in enumerate(sections, 1):
        _h(f"{i}. {section['title']}", 22, space_before=0, space_after=10)

        if section.get('description'):
            dp = doc.add_paragraph(style='Normal')
            dp.paragraph_format.alignment  = WD_ALIGN_PARAGRAPH.JUSTIFY
            dp.paragraph_format.space_after = Pt(12)
            dr = dp.add_run(section['description'])
            dr.italic = True
            dr.font.size = Pt(11)
            dr.font.color.rgb = RGBColor(0, 0, 0)

        for chapter in section.get('chapters', []):
            _h(chapter['title'], 14, space_before=12, space_after=6)
            content = chapter.get('content', '')
            if content:
                _render_html_docx(doc, content)

        # Page break between sections but not after the last one
        if i < len(sections):
            doc.add_page_break()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    response = make_response(buf.read())
    response.headers['Content-Type'] = (
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response.headers['Content-Disposition'] = 'attachment; filename="cos-mm-handbook.docx"'
    return response

@app.route("/admin/handbook/new", methods=["GET", "POST"])
@login_required
def admin_handbook_new():
    if request.method == "POST":
        max_order = get_max_sort_order("handbook_sections")
        with get_db().cursor() as c:
            c.execute(
                "INSERT INTO handbook_sections (id, title, description, sort_order) VALUES (%s, %s, %s, %s)",
                (new_id(), request.form["title"], request.form.get("description", ""), max_order + 1)
            )
        flash("Section added.", "success")
        return redirect(url_for("admin_handbook"))
    return render_template("admin/handbook_section_form.html", item=None)

@app.route("/admin/handbook/edit/<int:section_id>", methods=["GET", "POST"])
@login_required
def admin_handbook_edit(section_id):
    with get_db().cursor(cursor_factory=psycopg2.extras.RealDictCursor) as c:
        c.execute("SELECT * FROM handbook_sections WHERE id = %s", (section_id,))
        item = c.fetchone()
    if not item:
        flash("Not found.", "error")
        return redirect(url_for("admin_handbook"))
    item = dict(item)
    if request.method == "POST":
        with get_db().cursor() as c:
            c.execute(
                "UPDATE handbook_sections SET title=%s, description=%s WHERE id=%s",
                (request.form["title"], request.form.get("description", ""), section_id)
            )
        flash("Section updated.", "success")
        return redirect(url_for("admin_handbook"))
    return render_template("admin/handbook_section_form.html", item=item)

@app.route("/admin/handbook/delete/<int:section_id>", methods=["POST"])
@login_required
def admin_handbook_delete(section_id):
    with get_db().cursor() as c:
        c.execute("DELETE FROM handbook_sections WHERE id = %s", (section_id,))
    flash("Section deleted.", "success")
    return redirect(url_for("admin_handbook"))

@app.route("/admin/handbook/<int:section_id>/move/<direction>", methods=["POST"])
@login_required
def admin_handbook_section_move(section_id, direction):
    with get_db().cursor() as c:
        c.execute("SELECT id FROM handbook_sections ORDER BY sort_order, id")
        ids = [r[0] for r in c.fetchall()]
    idx = next((i for i, id_ in enumerate(ids) if id_ == section_id), None)
    if idx is not None:
        if direction == "up" and idx > 0:
            _swap_sort_order("handbook_sections", "id", ids[idx], ids[idx - 1])
        elif direction == "down" and idx < len(ids) - 1:
            _swap_sort_order("handbook_sections", "id", ids[idx], ids[idx + 1])
    return redirect(url_for("admin_handbook"))

# ── Admin: handbook chapters ──────────────────────────────────────────────────

@app.route("/admin/handbook/<int:section_id>/chapters")
@login_required
def admin_handbook_chapters(section_id):
    sections = load_handbook_meta()
    section = next((s for s in sections if s["id"] == section_id), None)
    if not section:
        flash("Section not found.", "error")
        return redirect(url_for("admin_handbook"))
    return render_template("admin/handbook_chapters.html", section=section)

@app.route("/admin/handbook/<int:section_id>/chapters/new", methods=["GET", "POST"])
@login_required
def admin_handbook_chapter_new(section_id):
    sections = load_handbook_meta()
    section = next((s for s in sections if s["id"] == section_id), None)
    if not section:
        return redirect(url_for("admin_handbook"))
    if request.method == "POST":
        max_order = get_max_sort_order("handbook_chapters", "section_id", section_id)
        with get_db().cursor() as c:
            c.execute(
                "INSERT INTO handbook_chapters (id, section_id, title, content, sort_order) VALUES (%s, %s, %s, %s, %s)",
                (new_id(), section_id, request.form["title"], request.form["content"], max_order + 1)
            )
        flash("Chapter added.", "success")
        return redirect(url_for("admin_handbook_chapters", section_id=section_id))
    return render_template("admin/handbook_chapter_form.html", section=section, item=None)

@app.route("/admin/handbook/<int:section_id>/chapters/edit/<int:chapter_id>", methods=["GET", "POST"])
@login_required
def admin_handbook_chapter_edit(section_id, chapter_id):
    # Fetch only this section and chapter — no need to load all content
    with get_db().cursor(cursor_factory=psycopg2.extras.RealDictCursor) as c:
        c.execute("SELECT id, title FROM handbook_sections WHERE id = %s", (section_id,))
        section_row = c.fetchone()
        if not section_row:
            return redirect(url_for("admin_handbook"))
        section = dict(section_row)
        c.execute(
            "SELECT id, title, content FROM handbook_chapters WHERE id = %s AND section_id = %s",
            (chapter_id, section_id)
        )
        chapter_row = c.fetchone()
        if not chapter_row:
            flash("Chapter not found.", "error")
            return redirect(url_for("admin_handbook_chapters", section_id=section_id))
        chapter = dict(chapter_row)
    if request.method == "POST":
        with get_db().cursor() as c:
            c.execute(
                "UPDATE handbook_chapters SET title=%s, content=%s WHERE id=%s",
                (request.form["title"], request.form["content"], chapter_id)
            )
        flash("Chapter updated.", "success")
        return redirect(url_for("admin_handbook_chapters", section_id=section_id))
    return render_template("admin/handbook_chapter_form.html", section=section, item=chapter)

@app.route("/admin/handbook/<int:section_id>/chapters/delete/<int:chapter_id>", methods=["POST"])
@login_required
def admin_handbook_chapter_delete(section_id, chapter_id):
    with get_db().cursor() as c:
        c.execute("DELETE FROM handbook_chapters WHERE id = %s AND section_id = %s", (chapter_id, section_id))
    flash("Chapter deleted.", "success")
    return redirect(url_for("admin_handbook_chapters", section_id=section_id))

@app.route("/admin/handbook/<int:section_id>/chapters/move/<int:chapter_id>/<direction>", methods=["POST"])
@login_required
def admin_handbook_chapter_move(section_id, chapter_id, direction):
    with get_db().cursor() as c:
        c.execute(
            "SELECT id FROM handbook_chapters WHERE section_id = %s ORDER BY sort_order, id",
            (section_id,)
        )
        ids = [r[0] for r in c.fetchall()]
    idx = next((i for i, id_ in enumerate(ids) if id_ == chapter_id), None)
    if idx is not None:
        if direction == "up" and idx > 0:
            _swap_sort_order("handbook_chapters", "id", ids[idx], ids[idx - 1])
        elif direction == "down" and idx < len(ids) - 1:
            _swap_sort_order("handbook_chapters", "id", ids[idx], ids[idx + 1])
    return redirect(url_for("admin_handbook_chapters", section_id=section_id))

# ── Admin: resources ──────────────────────────────────────────────────────────

@app.route("/admin/resources")
@login_required
def admin_resources():
    return render_template("admin/resources.html", categories=load_resources())

@app.route("/admin/resources/new-category", methods=["GET", "POST"])
@login_required
def admin_resource_new_category():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        if name:
            max_order = get_max_sort_order("resource_categories")
            try:
                with get_db().cursor() as c:
                    c.execute(
                        "INSERT INTO resource_categories (name, sort_order) VALUES (%s, %s)",
                        (name, max_order + 1)
                    )
                flash("Category added.", "success")
            except psycopg2.errors.UniqueViolation:
                get_db().rollback()
                flash("A category with that name already exists.", "error")
        return redirect(url_for("admin_resources"))
    return render_template("admin/resource_category_form.html")

@app.route("/admin/resources/new", methods=["GET", "POST"])
@login_required
def admin_resource_new():
    categories = load_resources()
    if request.method == "POST":
        cat_name = request.form["category"]
        cat_id = get_category_id_by_name(cat_name)
        if not cat_id:
            max_order = get_max_sort_order("resource_categories")
            with get_db().cursor() as c:
                c.execute(
                    "INSERT INTO resource_categories (name, sort_order) VALUES (%s, %s) RETURNING id",
                    (cat_name, max_order + 1)
                )
                cat_id = c.fetchone()[0]
        max_order = get_max_sort_order("resource_items", "category_id", cat_id)
        with get_db().cursor() as c:
            c.execute(
                """INSERT INTO resource_items
                   (id, category_id, title, description, url, content, type, sort_order)
                   VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
                (
                    new_id(), cat_id,
                    request.form["title"],
                    request.form.get("description", ""),
                    request.form.get("url", ""),
                    request.form.get("content", ""),
                    request.form.get("type", "link"),
                    max_order + 1,
                )
            )
        flash("Resource added.", "success")
        return redirect(url_for("admin_resources"))
    preset_category = request.args.get("category", "")
    cat_names = [c["name"] for c in categories]
    return render_template("admin/resource_form.html", item=None, categories=cat_names,
                           preset_category=preset_category)

@app.route("/admin/resources/edit/<int:item_id>", methods=["GET", "POST"])
@login_required
def admin_resource_edit(item_id):
    categories = load_resources()
    item = None
    item_category = None
    for cat in categories:
        for i in cat["items"]:
            if i["id"] == item_id:
                item = i
                item_category = cat["name"]
                break
        if item:
            break
    if not item:
        flash("Resource not found.", "error")
        return redirect(url_for("admin_resources"))
    if request.method == "POST":
        with get_db().cursor() as c:
            c.execute(
                "UPDATE resource_items SET title=%s, type=%s, url=%s, content=%s, description=%s WHERE id=%s",
                (
                    request.form["title"],
                    request.form.get("type", "link"),
                    request.form.get("url", ""),
                    request.form.get("content", ""),
                    request.form.get("description", ""),
                    item_id,
                )
            )
        flash("Resource updated.", "success")
        return redirect(url_for("admin_resources"))
    cat_names = [c["name"] for c in categories]
    return render_template("admin/resource_form.html", item=item, categories=cat_names,
                           preset_category=None, item_category=item_category)

@app.route("/admin/resources/move-category/<path:category_name>/<direction>", methods=["POST"])
@login_required
def admin_resource_category_move(category_name, direction):
    with get_db().cursor() as c:
        c.execute("SELECT id, name FROM resource_categories ORDER BY sort_order, id")
        cats = c.fetchall()
    idx = next((i for i, cat in enumerate(cats) if cat[1] == category_name), None)
    if idx is not None:
        if direction == "up" and idx > 0:
            _swap_sort_order("resource_categories", "id", cats[idx][0], cats[idx - 1][0])
        elif direction == "down" and idx < len(cats) - 1:
            _swap_sort_order("resource_categories", "id", cats[idx][0], cats[idx + 1][0])
    return redirect(url_for("admin_resources"))

@app.route("/admin/resources/rename-category", methods=["POST"])
@login_required
def admin_resource_rename_category():
    old_name = request.form.get("old_name", "").strip()
    new_name = request.form.get("new_name", "").strip()
    if old_name and new_name:
        with get_db().cursor() as c:
            c.execute(
                "UPDATE resource_categories SET name = %s WHERE name = %s",
                (new_name, old_name)
            )
        flash("Category renamed.", "success")
    return redirect(url_for("admin_resources"))

@app.route("/admin/resources/delete/<int:item_id>", methods=["POST"])
@login_required
def admin_resource_delete(item_id):
    with get_db().cursor() as c:
        c.execute("DELETE FROM resource_items WHERE id = %s", (item_id,))
    flash("Resource deleted.", "success")
    return redirect(url_for("admin_resources"))

# ── Run ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    app.run(debug=True, port=5003)
